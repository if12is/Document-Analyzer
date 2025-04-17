import os
import argparse
import time
import json
import re
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Configuration ---
load_dotenv()
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
GEMINI_MODEL_NAME = os.getenv('GEMINI_MODEL', 'gemini-2.0-flash')
print(f"[*] Using model: {GEMINI_MODEL_NAME}")
if not GOOGLE_API_KEY:
    print("Error: GOOGLE_API_KEY environment variable not found.")
    print("Please create a .env file with GOOGLE_API_KEY='Your-API-Key-Here'")
    exit()

try:
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel(GEMINI_MODEL_NAME)
    print(f"[*] Using Gemini model: {GEMINI_MODEL_NAME}")
    if "gemini-1.5" not in GEMINI_MODEL_NAME and "gemini-2" not in GEMINI_MODEL_NAME:
         print(f"[!] Warning: Model '{GEMINI_MODEL_NAME}' might not support direct file input or complex generation needed. "
               "Consider using 'gemini-1.5-flash-latest' or 'gemini-1.5-pro-latest'.")
except Exception as e:
    print(f"An unexpected error occurred during Gemini configuration: {e}")
    exit()

# --- Constants for Parsing ---
TEXT_START_MARKER = "--- START OF EXTRACTED TEXT ---"
TEXT_END_MARKER = "--- END OF EXTRACTED TEXT ---"
SUMMARY_START_MARKER = "--- START OF SUMMARY ---"
SUMMARY_END_MARKER = "--- END OF SUMMARY ---"

# --- Functions ---

def upload_file_to_gemini(file_path, retries=3, delay=5):
    """
    Uploads a file (PDF or image) to Google AI for processing with Gemini.
    Includes retry logic and waits for the file to become active.
    """
    print(f"[*] Uploading file: {file_path}...")
    if not os.path.exists(file_path):
        print(f"[!] Error: File not found at '{file_path}'")
        return None

    attempt = 0
    while attempt < retries:
        try:
            uploaded_file = genai.upload_file(path=file_path)
            print(f"[+] File uploaded successfully. File Name: {uploaded_file.name}")
            print("[*] Waiting for file processing...")
            processing_timeout = 300 # 5 minutes timeout for processing
            start_time = time.time()
            while uploaded_file.state.name == "PROCESSING":
                 if time.time() - start_time > processing_timeout:
                      print("[!] Error: File processing timed out.")
                      try:
                           genai.delete_file(uploaded_file.name)
                           print(f"[*] Cleaned up timed-out file: {uploaded_file.name}")
                      except Exception as delete_e:
                           print(f"[!] Warning: Could not delete timed-out file {uploaded_file.name}: {delete_e}")
                      return None
                 time.sleep(15)
                 uploaded_file = genai.get_file(uploaded_file.name)
                 print(f"[*] Current file state: {uploaded_file.state.name}")

            if uploaded_file.state.name == "ACTIVE":
                 print("[+] File is active and ready.")
                 return uploaded_file
            else:
                 print(f"[!] Error: File processing failed or ended in unexpected state: {uploaded_file.state.name}")
                 try:
                      if uploaded_file.state.name != "DELETED":
                           genai.delete_file(uploaded_file.name)
                           print(f"[*] Cleaned up file in state {uploaded_file.state.name}: {uploaded_file.name}")
                 except Exception as delete_e:
                      print(f"[!] Warning: Could not delete file {uploaded_file.name} in state {uploaded_file.state.name}: {delete_e}")
                 return None

        except Exception as e:
            attempt += 1
            print(f"[!] Error uploading file (Attempt {attempt}/{retries}): {e}")
            if attempt < retries:
                print(f"[*] Retrying in {delay} seconds...")
                time.sleep(delay)
            else:
                print("[!] Upload failed after multiple retries.")
                return None
    return None

def save_to_text_file(content, output_path):
    """
    Saves the given content to a text file using UTF-8 encoding.
    Creates parent directories if they don't exist.
    """
    print(f"[*] Saving content to text file: {output_path}")
    if content is None:
         print(f"[!] Cannot save text: content is None.")
         return False
    try:
        directory = os.path.dirname(output_path)
        if directory:
            os.makedirs(directory, exist_ok=True)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"[+] Content saved successfully to text file.")
        return True
    except Exception as e:
        print(f"[!] Error saving text file to '{output_path}': {e}")
        return False

def save_to_word_file(content, output_path, title, language="arabic"):
    """
    Saves the given content to a Word (.docx) file.
    Creates parent directories if they don't exist.
    """
    print(f"[*] Saving content to Word file: {output_path}")
    if content is None:
         print(f"[!] Cannot save Word document: content is None.")
         return False
    try:
        directory = os.path.dirname(output_path)
        if directory:
            os.makedirs(directory, exist_ok=True)

        doc = Document()
        
        # Configure direction based on language
        if language == "arabic":
            # Handle RTL for Arabic
            doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rtl = True
        else:
            rtl = False

        # Add a title
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        # Add a line break
        doc.add_paragraph()
        
        # Add the content
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph()
                if rtl:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.add_run(para)
        
        # Save the document
        doc.save(output_path)
        print(f"[+] Content saved successfully to Word document.")
        return True
    except Exception as e:
        print(f"[!] Error saving Word file to '{output_path}': {e}")
        return False

def analyze_document(uploaded_file, extract_mode="full", language="arabic"):
    """
    Analyzes a document (PDF or image) and extracts text content.
    
    Args:
        uploaded_file: The uploaded file object from Gemini
        extract_mode: 'full' for full text extraction or 'summary' for summarized content
        language: 'arabic' or 'english' for output language
    
    Returns:
        tuple: (extracted_text, summary_text) - Either or both may be None if extraction fails
    """
    print(f"[*] Analyzing document in {language}, mode: {extract_mode}...")

    # Determine text direction based on language for prompt clarity
    text_dir = "right-to-left" if language == "arabic" else "left-to-right"
    
    # Construct the prompt based on extraction mode
    if extract_mode == "summary":
        prompt = f"""
Please analyze the attached document ({uploaded_file.display_name}) and provide the following in {language}:

**Task 1: Extract and Summarize Content**
1. Extract all visible text from the document (including text in images if possible).
2. Create a comprehensive summary of the document's content.
3. Include key points, main ideas, and important details.
4. Format the summary in clear, well-structured paragraphs.
5. The summary should be approximately 20-30% of the original text length.
6. Ensure the text follows {text_dir} direction appropriate for {language}.

**Task 2: Full Text Extraction**
1. Extract ALL textual content from the document, including any text visible in images.
2. Present the extracted text in a clean, readable format.
3. Preserve paragraph breaks and important formatting where possible.
4. Ensure the text follows {text_dir} direction appropriate for {language}.

**Output Format:**
Please structure your response exactly as follows:

{SUMMARY_START_MARKER}
[Your detailed summary here]
{SUMMARY_END_MARKER}

{TEXT_START_MARKER}
[Full extracted text here]
{TEXT_END_MARKER}
"""
    else:  # full extraction mode
        prompt = f"""
Please analyze the attached document ({uploaded_file.display_name}) and provide the following in {language}:

**Task: Full Text Extraction**
1. Extract ALL textual content from the document, including any text visible in images.
2. Present the extracted text in a clean, readable format.
3. Preserve paragraph breaks and important formatting where possible.
4. Ensure the text follows {text_dir} direction appropriate for {language}.
5. DO NOT summarize or alter the content - provide the complete extracted text only.

**Output Format:**
Please structure your response exactly as follows:

{TEXT_START_MARKER}
[Full extracted text here]
{TEXT_END_MARKER}
"""

    extracted_text = None
    summary_text = None

    try:
        # Call the Gemini API
        print("[*] Sending request to Gemini API...")
        request_options = {"timeout": 600.0} # 10 minutes timeout

        response = model.generate_content(
            [prompt, uploaded_file],
            request_options=request_options
            )

        # Parse the response
        print("[*] Parsing Gemini response...")
        if not hasattr(response, 'text'):
             print("[!] Error: Gemini response did not contain expected text content.")
             if hasattr(response, 'prompt_feedback'):
                  print(f"[*] Prompt Feedback: {response.prompt_feedback}")
             return None, None

        full_response_text = response.text

        # Extract full text
        text_match = re.search(f"{re.escape(TEXT_START_MARKER)}(.*?){re.escape(TEXT_END_MARKER)}", 
                              full_response_text, re.DOTALL | re.IGNORECASE)
        if text_match:
            extracted_text = text_match.group(1).strip()
            print("[+] Full text content extracted successfully.")
        else:
            print(f"[!] Warning: Could not find extracted text between markers.")

        # Extract summary if in summary mode
        if extract_mode == "summary":
            summary_match = re.search(f"{re.escape(SUMMARY_START_MARKER)}(.*?){re.escape(SUMMARY_END_MARKER)}", 
                                     full_response_text, re.DOTALL | re.IGNORECASE)
            if summary_match:
                summary_text = summary_match.group(1).strip()
                print("[+] Summary content extracted successfully.")
            else:
                print(f"[!] Warning: Could not find summary between markers.")

        # If we didn't extract anything, check if there's any text in the response
        if extracted_text is None and summary_text is None:
            print("[!] Warning: Failed to parse both text and summary using markers.")
            if full_response_text.strip():
                extracted_text = full_response_text.strip()
                print("[*] Using full response as extracted text.")

        return extracted_text, summary_text

    except Exception as e:
        print(f"[!] An error occurred during document analysis: {e}")
        return None, None
    finally:
        # Clean up the uploaded file
        try:
            if uploaded_file and uploaded_file.name:
                print(f"[*] Deleting uploaded file from cloud storage: {uploaded_file.name}")
                genai.delete_file(uploaded_file.name)
                print("[+] File deleted successfully.")
        except Exception as delete_e:
            print(f"[!] Warning: Failed to delete uploaded file. Error: {delete_e}")


# --- Main Execution ---
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract text from documents (PDFs or images) using Gemini AI.")

    parser.add_argument("input_file", help="Path to the input file (PDF or image).")
    parser.add_argument("-o", "--output_file", help="Path to save the extracted text. Default is [filename]_extracted.[txt/docx]")
    parser.add_argument("-f", "--format", choices=['text', 'docx'], default='text', 
                        help="Output format: text (.txt) or Word document (.docx). Default is text.")
    parser.add_argument("-m", "--mode", choices=['full', 'summary'], default='full',
                        help="Extraction mode: full text or summary. Default is full.")
    parser.add_argument("-l", "--language", choices=['arabic', 'english'], default='arabic',
                        help="Language for the output. Default is arabic.")

    args = parser.parse_args()

    # Set default output filename if not provided
    base_filename = os.path.splitext(os.path.basename(args.input_file))[0]
    default_ext = ".docx" if args.format == "docx" else ".txt"
    output_path = args.output_file if args.output_file else f"{base_filename}_{args.mode}_{args.language}{default_ext}"

    # 1. Upload file
    uploaded_file = upload_file_to_gemini(args.input_file)

    if uploaded_file:
        # 2. Analyze document
        extracted_text, summary_text = analyze_document(uploaded_file, args.mode, args.language)

        # 3. Save results
        content_to_save = summary_text if args.mode == "summary" and summary_text else extracted_text
        
        if content_to_save:
            if args.format == "docx":
                title = f"{'Summary' if args.mode == 'summary' else 'Extracted Text'} - {base_filename}"
                success = save_to_word_file(content_to_save, output_path, title, args.language)
            else:
                success = save_to_text_file(content_to_save, output_path)
                
            if success:
                print(f"\n[+] Content saved to: {output_path}")
            else:
                print("\n[!] Failed to save content.")
        else:
            print("\n[!] No content to save. Extraction failed.")
    else:
        print("\n[!] Process failed: Could not upload or process the file.")

    print("[*] Script finished.")
